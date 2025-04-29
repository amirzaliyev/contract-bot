from __future__ import annotations
from pathlib import Path
from typing import Dict

from docx import Document as PyDocument
from uuid import uuid4

from config import BOT_DIR, DOCS_DIR
from data.models import Document
from data.repositories.document_repository import DocumentNotFound, IDocumentRepository
    

class DocumentCreationUnsuccessful(Exception):
    pass    

    
class DocumentService():
    """
    Creates documents from the templates
    """
    
    def __init__(self, doc_repository: IDocumentRepository) -> None:
        
        if not isinstance(doc_repository, IDocumentRepository):
            raise TypeError(
                "doc_repository should be an instance of 'IDocumentRepository' not " +
                f"{type(doc_repository).__name__}"
            )
            
        self._repo = doc_repository
        
    def _get_template_path(self, user_id: int, template_id: int) -> str:
        try:
            template = self._repo.get_document_by_id(
                owner_id=user_id, 
                document_id=template_id, 
                is_template=True
            )
        
        except DocumentNotFound as e:
            raise DocumentCreationUnsuccessful(f"Document is not created because {str(e).lower()}")
        
        template_path = Path.joinpath(BOT_DIR, template.file_path)
        
        return str(template_path)
        
        
    def fill_template(self, user_id: int, template_id: int, context: Dict[str, str]) -> str:
        template_path = self._get_template_path(user_id=user_id, template_id=template_id)
        
        file_name = self._generate_file_name()
        output_path = Path.joinpath(DOCS_DIR, f"{user_id}", file_name)
        relative_path = str(output_path.relative_to(BOT_DIR).as_posix())
        
        # file path as string
        output_path_str = str(output_path)
        
        # save to local docs folder
        self._save_to_file(
            template_path=template_path,
            output_path=output_path_str,
            context=context
        )
        
        # save to db
        self._save_to_db(
            document=Document(
                file_name=file_name,
                owner_id=user_id,
                file_path=relative_path
            )
        )
       
        
        return output_path_str
    
    
    def _save_to_file(self, template_path, output_path, context: Dict[str, str]) -> None:
        doc = PyDocument(template_path)

        for paragraph in doc.paragraphs:
            for key, value in context.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context.items():
                        if f"{{{{{key}}}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))
          
                            
        doc.save(output_path)
    
    
    def _save_to_db(self, document: Document) -> None:
        self._repo.create_document(document=document)
        
    
    def _generate_file_name(self, extension='docx') -> str:
        return f"{uuid4()}.{extension}"
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
     
        

# # === Step 1: Load data from Excel ===
# excel_path = "clients.xlsx"  # <-- update this to your actual Excel file name
# df = pd.read_excel(excel_path)

# # === Step 2: Setup templates and output folder ===
# template1 = "template1.docx"
# template2 = "template2.docx"
# output_folder_c = "shartnomalar"
# output_folder_o = "buyruqlar"
# os.makedirs(output_folder_c, exist_ok=True)
# os.makedirs(output_folder_o, exist_ok=True)

# # === Step 3: Generate documents for each row ===
# for index, row in df.iterrows():
#     context = row.to_dict()

#     name_for_file = context["full_name"].replace(" ", "_")
#     print(context)
#     doc_id = context["document_id"]

#     output_path1 = os.path.join(output_folder_c, f"Mehnat_Shartnoma_{name_for_file}.docx")
#     output_path2 = os.path.join(output_folder_o, f"Buyruq SHT-{'0' if doc_id - 1 < 10 else ''}{doc_id - 1}.docx")

#     fill_template(template1, output_path1, context)
#     fill_template(template2, output_path2, context)

# print("All documents generated successfully.")
